#!/usr/bin/env python3
"""
Excel to Word Table Extractor
Extracts a specified block of data from an Excel file and creates a Word document with a table.
"""

import sys
import os
from pathlib import Path
import openpyxl
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import configuration
try:
    import config
except ImportError:
    print("Error: config.py not found. Please ensure config.py exists in the same directory.")
    sys.exit(1)


def extract_excel_data(excel_file, sheet_name, start_row, end_row, start_col, end_col):
    """
    Extract data from Excel file within specified range.
    
    Args:
        excel_file (str): Path to Excel file
        sheet_name (str): Name of the worksheet (or None for active sheet)
        start_row (int): Starting row (1-indexed)
        end_row (int): Ending row (1-indexed, inclusive)
        start_col (int): Starting column (1-indexed)
        end_col (int): Ending column (1-indexed, inclusive)
    
    Returns:
        list: 2D list containing the extracted data
    """
    try:
        print(f"Opening Excel file: {excel_file}")
        workbook = openpyxl.load_workbook(excel_file, data_only=True)
        
        # Use specified sheet or active sheet
        if sheet_name:
            if sheet_name not in workbook.sheetnames:
                print(f"Error: Sheet '{sheet_name}' not found.")
                print(f"Available sheets: {', '.join(workbook.sheetnames)}")
                return None
            worksheet = workbook[sheet_name]
            print(f"Using sheet: {sheet_name}")
        else:
            worksheet = workbook.active
            print(f"Using active sheet: {worksheet.title}")
        
        # Validate range
        max_row = worksheet.max_row
        max_col = worksheet.max_column
        
        if start_row > max_row or end_row > max_row:
            print(f"Warning: Specified rows exceed sheet maximum ({max_row})")
            end_row = min(end_row, max_row)
        
        if start_col > max_col or end_col > max_col:
            print(f"Warning: Specified columns exceed sheet maximum ({max_col})")
            end_col = min(end_col, max_col)
        
        # Extract data from specified range
        data = []
        for row in range(start_row, end_row + 1):
            row_data = []
            for col in range(start_col, end_col + 1):
                cell_value = worksheet.cell(row=row, column=col).value
                # Convert None to empty string for better display
                row_data.append(str(cell_value) if cell_value is not None else "")
            data.append(row_data)
        
        workbook.close()
        return data
        
    except FileNotFoundError:
        print(f"Error: Excel file '{excel_file}' not found.")
        return None
    except PermissionError:
        print(f"Error: Permission denied to read '{excel_file}'.")
        print("Please ensure the file is not open in another program.")
        return None
    except Exception as e:
        print(f"Error reading Excel file: {str(e)}")
        return None


def create_word_table(data, output_file, title, config):
    """
    Create a Word document with a table containing the extracted data.
    
    Args:
        data (list): 2D list containing the data
        output_file (str): Path for output Word document
        title (str): Title for the document
        config: Configuration module with formatting options
    
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        # Create new Word document
        doc = Document()
        
        # Add title
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata paragraph
        info_para = doc.add_paragraph()
        info_para.add_run(f"Source: {config.EXCEL_FILE}\n").italic = True
        info_para.add_run(f"Range: Rows {config.START_ROW}-{config.END_ROW}, ").italic = True
        info_para.add_run(f"Columns {config.START_COL}-{config.END_COL}\n").italic = True
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some spacing
        doc.add_paragraph()
        
        # Create table with dimensions matching data
        rows = len(data)
        cols = len(data[0]) if data else 0
        
        if rows == 0 or cols == 0:
            print("No data to export.")
            return False
        
        print(f"Creating table with {rows} rows and {cols} columns...")
        table = doc.add_table(rows=rows, cols=cols)
        table.style = config.TABLE_STYLE
        
        if config.CENTER_TABLE:
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Populate table with data
        for i, row_data in enumerate(data):
            for j, cell_data in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = str(cell_data)
                
                # Format header row if specified
                if config.FIRST_ROW_IS_HEADER and i == 0:
                    # Make header row bold
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                    # Optional: Add background color to header
                    cell._element.get_or_add_tcPr().append(
                        parse_xml(r'<w:shd {} w:fill="D3D3D3"/>'.format(
                            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                        ))
                    )
        
        # Auto-fit table width if specified
        if config.AUTO_FIT:
            table.autofit = True
        
        # Add footer with timestamp
        from datetime import datetime
        footer_para = doc.add_paragraph()
        footer_para.add_run(f"\nGenerated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Save document
        doc.save(output_file)
        print(f"✓ Word document saved as: {output_file}")
        return True
        
    except PermissionError:
        print(f"Error: Permission denied to write '{output_file}'.")
        print("Please ensure the file is not open in another program.")
        return False
    except Exception as e:
        print(f"Error creating Word document: {str(e)}")
        return False


def column_letter_to_number(column_letter):
    """
    Convert Excel column letter(s) to column number.
    A -> 1, B -> 2, Z -> 26, AA -> 27, etc.
    """
    column_number = 0
    for char in column_letter.upper():
        column_number = column_number * 26 + (ord(char) - ord('A') + 1)
    return column_number


def parse_xml(xml_string):
    """Helper function to parse XML for table cell formatting."""
    from docx.oxml import parse_xml as docx_parse_xml
    return docx_parse_xml(xml_string)


def validate_config():
    """
    Validate configuration settings.
    
    Returns:
        bool: True if valid, False otherwise
    """
    errors = []
    
    # Check if Excel file exists
    if not os.path.exists(config.EXCEL_FILE):
        errors.append(f"Excel file '{config.EXCEL_FILE}' not found")
    
    # Validate row numbers
    if config.START_ROW < 1:
        errors.append("START_ROW must be at least 1")
    if config.END_ROW < config.START_ROW:
        errors.append("END_ROW must be greater than or equal to START_ROW")
    
    # Validate column letters
    try:
        start_col_num = column_letter_to_number(config.START_COL)
        end_col_num = column_letter_to_number(config.END_COL)
        if start_col_num < 1:
            errors.append("START_COL must be a valid column letter (A, B, C, etc.)")
        if end_col_num < start_col_num:
            errors.append(f"END_COL ({config.END_COL}) must be after or equal to START_COL ({config.START_COL})")
    except (AttributeError, ValueError):
        errors.append("START_COL and END_COL must be valid column letters (A, B, C, etc.)")
    
    # Check output file path and create directory if needed
    output_dir = os.path.dirname(config.OUTPUT_FILE) or '.'
    if output_dir != '.' and not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
            print(f"✓ Created output directory: {output_dir}")
        except PermissionError:
            errors.append(f"Permission denied to create output directory '{output_dir}'")
        except Exception as e:
            errors.append(f"Failed to create output directory '{output_dir}': {str(e)}")
    
    if errors:
        print("Configuration errors:")
        for error in errors:
            print(f"  - {error}")
        return False
    
    return True


def main():
    """Main function to coordinate the extraction and export process."""
    print("=" * 50)
    print("Excel to Word Table Extractor")
    print("=" * 50)
    print(f"Input file: {config.EXCEL_FILE}")
    print(f"Sheet: {'Active sheet' if config.SHEET_NAME is None else config.SHEET_NAME}")
    print(f"Range: Rows {config.START_ROW}-{config.END_ROW}, Columns {config.START_COL}-{config.END_COL}")
    print(f"Output: {config.OUTPUT_FILE}")
    print("-" * 50)
    
    # Validate configuration
    if not validate_config():
        print("\n✗ Please fix configuration errors in config.py and try again.")
        return 1
    
    # Convert column letters to numbers
    start_col_num = column_letter_to_number(config.START_COL)
    end_col_num = column_letter_to_number(config.END_COL)
    
    # Extract data from Excel
    print("\nExtracting data from Excel...")
    data = extract_excel_data(
        config.EXCEL_FILE, 
        config.SHEET_NAME,
        config.START_ROW, 
        config.END_ROW,
        start_col_num, 
        end_col_num
    )
    
    if data is None:
        print("\n✗ Failed to extract data from Excel file.")
        return 1
    
    print(f"✓ Successfully extracted {len(data)} rows with {len(data[0]) if data else 0} columns.")
    
    # Preview first few rows
    if data:
        print("\nPreview of extracted data (first 5 rows):")
        print("-" * 50)
        for i, row in enumerate(data[:5]):
            preview = " | ".join(cell[:20] + "..." if len(cell) > 20 else cell for cell in row)
            print(f"Row {i+1}: {preview}")
        if len(data) > 5:
            print(f"... and {len(data) - 5} more rows")
    
    # Create Word document
    print("\nCreating Word document...")
    success = create_word_table(data, config.OUTPUT_FILE, config.DOCUMENT_TITLE, config)
    
    if success:
        print("\n✓ Export completed successfully!")
        print(f"Output file: {os.path.abspath(config.OUTPUT_FILE)}")
        return 0
    else:
        print("\n✗ Export failed.")
        return 1


if __name__ == "__main__":
    try:
        sys.exit(main())
    except KeyboardInterrupt:
        print("\n\nOperation cancelled by user.")
        sys.exit(1)
    except Exception as e:
        print(f"\nUnexpected error: {str(e)}")
        sys.exit(1)